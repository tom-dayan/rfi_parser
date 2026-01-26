import PyPDF2
import pdfplumber
from docx import Document
from typing import Optional
import io


class DocumentParser:
    """Parse documents (PDF, DOCX, TXT/MD) and extract text content"""

    @staticmethod
    def parse_pdf(file_content: bytes) -> str:
        """
        Parse PDF document using PyPDF2 and pdfplumber as fallback

        Args:
            file_content: Raw bytes of PDF file

        Returns:
            Extracted text content
        """
        text_content = []

        # Try PyPDF2 first (faster)
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text:
                    text_content.append(text)
        except Exception as e:
            print(f"PyPDF2 parsing failed: {e}, trying pdfplumber")

            # Fallback to pdfplumber (better for complex layouts)
            try:
                pdf_file = io.BytesIO(file_content)
                with pdfplumber.open(pdf_file) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            text_content.append(text)
            except Exception as e:
                raise ValueError(f"Failed to parse PDF: {e}")

        return "\n\n".join(text_content)

    @staticmethod
    def parse_docx(file_content: bytes) -> str:
        """
        Parse DOCX document

        Args:
            file_content: Raw bytes of DOCX file

        Returns:
            Extracted text content
        """
        try:
            doc_file = io.BytesIO(file_content)
            doc = Document(doc_file)

            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Extract table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_content.append(row_text)

            return "\n\n".join(text_content)
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {e}")

    @staticmethod
    def parse_text(file_content: bytes) -> str:
        """
        Parse plain text or markdown document

        Args:
            file_content: Raw bytes of text file

        Returns:
            Decoded text content
        """
        try:
            # Try UTF-8 first
            return file_content.decode('utf-8')
        except UnicodeDecodeError:
            # Fallback to latin-1
            try:
                return file_content.decode('latin-1')
            except Exception as e:
                raise ValueError(f"Failed to decode text file: {e}")

    @classmethod
    def parse_document(cls, file_content: bytes, filename: str) -> str:
        """
        Parse document based on file extension

        Args:
            file_content: Raw bytes of file
            filename: Name of the file (used to determine type)

        Returns:
            Extracted text content
        """
        filename_lower = filename.lower()

        if filename_lower.endswith('.pdf'):
            return cls.parse_pdf(file_content)
        elif filename_lower.endswith('.docx') or filename_lower.endswith('.doc'):
            return cls.parse_docx(file_content)
        elif filename_lower.endswith(('.txt', '.md', '.markdown')):
            return cls.parse_text(file_content)
        else:
            # Try to parse as text by default
            return cls.parse_text(file_content)

    @staticmethod
    def extract_sections(content: str) -> dict[str, str]:
        """
        Extract logical sections from document content

        Args:
            content: Full document text

        Returns:
            Dictionary of section titles to content
        """
        sections = {}
        current_section = "Introduction"
        current_content = []

        lines = content.split('\n')

        for line in lines:
            # Simple heuristic: lines that are short, uppercase, or end with numbers might be headers
            stripped = line.strip()

            # Detect section headers (simple heuristic)
            is_header = (
                (stripped and len(stripped) < 100 and
                 (stripped.isupper() or
                  any(stripped.startswith(prefix) for prefix in ['Section', 'Chapter', 'Article', 'Clause', 'SECTION', 'CHAPTER'])))
            )

            if is_header:
                # Save previous section
                if current_content:
                    sections[current_section] = '\n'.join(current_content).strip()

                # Start new section
                current_section = stripped
                current_content = []
            else:
                current_content.append(line)

        # Save last section
        if current_content:
            sections[current_section] = '\n'.join(current_content).strip()

        return sections
