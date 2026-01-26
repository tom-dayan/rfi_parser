import io
from docx import Document
from .base import DocumentParser, ParseResult


class DocxParser(DocumentParser):
    """Parser for Microsoft Word documents"""

    supported_extensions = ['docx', 'doc']

    def parse(self, file_path: str) -> ParseResult:
        """Parse DOCX file from path"""
        try:
            content = self._read_file(file_path)
            return self.parse_bytes(content, file_path)
        except Exception as e:
            return ParseResult.error_result(f"Failed to read file: {str(e)}")

    def parse_bytes(self, content: bytes, filename: str) -> ParseResult:
        """Parse DOCX from bytes"""
        try:
            doc_file = io.BytesIO(content)
            doc = Document(doc_file)

            text_parts = []
            metadata = {
                'paragraph_count': 0,
                'table_count': 0,
                'has_images': False,
            }

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text and para.text.strip():
                    text_parts.append(para.text)
                    metadata['paragraph_count'] += 1

            # Extract tables
            for table in doc.tables:
                metadata['table_count'] += 1
                table_text = self._extract_table(table)
                if table_text:
                    text_parts.append(f"[Table]\n{table_text}")

            # Check for images (inline shapes)
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    metadata['has_images'] = True
                    break

            # Extract document properties if available
            core_props = doc.core_properties
            if core_props.title:
                metadata['title'] = core_props.title
            if core_props.author:
                metadata['author'] = core_props.author
            if core_props.subject:
                metadata['subject'] = core_props.subject

            if not text_parts:
                return ParseResult.success_result(
                    text="[Empty or image-only document]",
                    metadata=metadata
                )

            return ParseResult.success_result(
                text="\n\n".join(text_parts),
                metadata=metadata
            )

        except Exception as e:
            return ParseResult.error_result(f"Failed to parse DOCX: {str(e)}")

    def _extract_table(self, table) -> str:
        """Extract text from a Word table"""
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip() if cell.text else ''
                cells.append(cell_text)
            if any(cells):
                rows.append(' | '.join(cells))
        return '\n'.join(rows) if rows else ''
